from docx import Document

def generate_docx(summary, skills, matched_keywords, missing_keywords):
    doc = Document()
    doc.add_heading("AI Resume Tailor Report", 0)

    # Summary
    doc.add_heading("🧠 Tailored Summary", level=1)
    doc.add_paragraph(summary if summary else "No summary available.")

    # Skills
    doc.add_heading("💼 Suggested Skills Section", level=1)
    if skills:
        for skill in skills:
            doc.add_paragraph(f"• {skill}")
    else:
        doc.add_paragraph("No skills generated.")

    # Matched Keywords
    doc.add_heading("✅ Matched Keywords", level=1)
    doc.add_paragraph(", ".join(sorted(matched_keywords)) if matched_keywords else "None")

    # Missing Keywords
    doc.add_heading("❌ Missing Important Keywords", level=1)
    doc.add_paragraph(", ".join(sorted(missing_keywords)) if missing_keywords else "None")

    # Save
    from io import BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
